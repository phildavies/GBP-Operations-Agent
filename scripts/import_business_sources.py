from __future__ import annotations

import argparse
import json
import mimetypes
import re
import sys
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = PROJECT_ROOT / "reports" / "generated"
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".tif", ".tiff"}
TEXT_EXTENSIONS = {".txt", ".md", ".html", ".htm", ".csv", ".json", ".rtf"}
DOCX_EXTENSIONS = {".docx"}
XLSX_EXTENSIONS = {".xlsx", ".xlsm"}
PDF_EXTENSIONS = {".pdf"}


@dataclass
class SourceResult:
    label: str
    path: str
    kind: str
    status: str
    summary: dict[str, Any]
    text_excerpt: str = ""
    error: str | None = None


def parse_manifest(path: Path) -> dict[str, Any]:
    lines = path.read_text(encoding="utf-8").splitlines()
    result: dict[str, Any] = {}
    section: str | None = None
    subsection: str | None = None

    for raw in lines:
        line = raw.rstrip()
        if not line.strip() or line.lstrip().startswith("#"):
            continue

        if not raw.startswith(" "):
            key, _, value = line.partition(":")
            key = key.strip()
            value = value.strip()
            subsection = None

            if value:
                result[key] = value
                section = None
            else:
                result[key] = {}
                section = key
            continue

        if section is None:
            continue

        stripped = line.strip()
        if stripped.startswith("- "):
            if subsection is not None:
                result[section].setdefault(subsection, []).append(stripped[2:].strip())
            else:
                if not isinstance(result.get(section), list):
                    result[section] = []
                result[section].append(stripped[2:].strip())
            continue

        key, _, value = stripped.partition(":")
        key = key.strip()
        value = value.strip()

        if value:
            result[section][key] = value
            subsection = None
        else:
            result[section][key] = []
            subsection = key

    return result


def resolve_source(source_root: Path, value: str) -> Path:
    value_path = Path(value)
    if value_path.is_absolute():
        return value_path
    return source_root / value


def compact_text(text: str, limit: int = 2500) -> str:
    normalized = re.sub(r"\s+", " ", text).strip()
    if len(normalized) <= limit:
        return normalized
    return normalized[:limit].rsplit(" ", 1)[0] + "..."


def redact_exclusions(text: str, exclusions: list[str]) -> str:
    redacted = re.sub(r"https?://(?:www\.)?paypal\.me/\S+", "[excluded obsolete payment link]", text, flags=re.IGNORECASE)
    for exclusion in exclusions:
        if not exclusion:
            continue
        redacted = re.sub(re.escape(exclusion), "[excluded obsolete reference]", redacted, flags=re.IGNORECASE)
    return redacted


def classify_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in IMAGE_EXTENSIONS:
        return "image"
    if suffix in DOCX_EXTENSIONS:
        return "document"
    if suffix in XLSX_EXTENSIONS:
        return "spreadsheet"
    if suffix in PDF_EXTENSIONS:
        return "pdf"
    if suffix in TEXT_EXTENSIONS:
        return "text"
    return mimetypes.guess_type(path.name)[0] or "file"


def safe_stat(path: Path) -> dict[str, Any]:
    stat = path.stat()
    return {
        "sizeBytes": stat.st_size,
        "modified": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat()
    }


def extract_docx(path: Path) -> tuple[dict[str, Any], str]:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows[:20]:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    text = "\n".join(paragraphs)
    summary = {
        **safe_stat(path),
        "paragraphCount": len(paragraphs),
        "tableCount": len(document.tables),
        "tablesSample": tables[:2]
    }
    return summary, compact_text(text)


def extract_xlsx(path: Path) -> tuple[dict[str, Any], str]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    sheets = []
    text_lines = []

    for sheet in workbook.worksheets:
        row_count = 0
        sample_rows = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                row_count += 1
                if len(sample_rows) < 15:
                    sample_rows.append(values)
                text_lines.append(" | ".join(values))

        sheets.append({
            "name": sheet.title,
            "rowsWithValues": row_count,
            "sampleRows": sample_rows
        })

    summary = {
        **safe_stat(path),
        "sheetCount": len(workbook.worksheets),
        "sheets": sheets
    }
    return summary, compact_text("\n".join(text_lines))


def extract_pdf(path: Path) -> tuple[dict[str, Any], str]:
    reader = PdfReader(str(path))
    text_parts = []
    for page in reader.pages[:10]:
        text_parts.append(page.extract_text() or "")

    summary = {
        **safe_stat(path),
        "pageCount": len(reader.pages),
        "sampledPages": min(len(reader.pages), 10)
    }
    return summary, compact_text("\n".join(text_parts))


def extract_text_file(path: Path) -> tuple[dict[str, Any], str]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    summary = {
        **safe_stat(path),
        "lineCount": len(text.splitlines())
    }
    return summary, compact_text(text)


def inspect_image(path: Path) -> tuple[dict[str, Any], str]:
    with Image.open(path) as image:
        width, height = image.size
        summary = {
            **safe_stat(path),
            "width": width,
            "height": height,
            "format": image.format,
            "mode": image.mode,
            "aspectRatio": round(width / height, 3) if height else None
        }
    return summary, ""


def inspect_file(label: str, path: Path, exclusions: list[str]) -> SourceResult:
    kind = classify_file(path)
    try:
        if not path.exists():
            return SourceResult(label, str(path), kind, "missing", {}, error="Path does not exist.")
        if path.is_dir():
            files = [child for child in path.rglob("*") if child.is_file()]
            image_count = sum(1 for child in files if child.suffix.lower() in IMAGE_EXTENSIONS)
            return SourceResult(
                label,
                str(path),
                "directory",
                "ok",
                {
                    "fileCount": len(files),
                    "imageCount": image_count,
                    "sampleFiles": [str(child.relative_to(path)) for child in files[:30]]
                }
            )

        suffix = path.suffix.lower()
        if suffix in DOCX_EXTENSIONS:
            summary, text = extract_docx(path)
        elif suffix in XLSX_EXTENSIONS:
            summary, text = extract_xlsx(path)
        elif suffix in PDF_EXTENSIONS:
            summary, text = extract_pdf(path)
        elif suffix in IMAGE_EXTENSIONS:
            summary, text = inspect_image(path)
        elif suffix in TEXT_EXTENSIONS:
            summary, text = extract_text_file(path)
        else:
            summary, text = safe_stat(path), ""

        return SourceResult(label, str(path), kind, "ok", summary, redact_exclusions(text, exclusions))
    except Exception as exc:
        return SourceResult(label, str(path), kind, "error", {}, error=str(exc))


def collect_sources(manifest: dict[str, Any]) -> list[tuple[str, str]]:
    sources: list[tuple[str, str]] = []
    for group_name in ("contentSources", "mediaSources"):
        group = manifest.get(group_name, {})
        if not isinstance(group, dict):
            continue
        for label, values in group.items():
            if isinstance(values, list):
                for value in values:
                    sources.append((f"{group_name}.{label}", value))
            elif isinstance(values, str):
                sources.append((f"{group_name}.{label}", values))
    return sources


def summarize_knowledge_pack(results: list[SourceResult]) -> dict[str, Any]:
    ok = [result for result in results if result.status == "ok"]
    errors = [result for result in results if result.status == "error"]
    missing = [result for result in results if result.status == "missing"]
    documents = [result for result in ok if result.kind in {"document", "spreadsheet", "pdf", "text"}]
    media = [result for result in ok if result.kind in {"image", "directory"}]

    extracted_text = "\n\n".join(
        f"## {Path(result.path).name}\n{result.text_excerpt}"
        for result in documents
        if result.text_excerpt
    )

    return {
        "sourceCount": len(results),
        "okCount": len(ok),
        "errorCount": len(errors),
        "missingCount": len(missing),
        "documentCount": len(documents),
        "mediaSourceCount": len(media),
        "combinedTextExcerpt": compact_text(extracted_text, 8000)
    }


def write_markdown(pack: dict[str, Any], output_path: Path) -> None:
    manifest = pack["manifest"]
    summary = pack["summary"]
    lines = [
        f"# Business Knowledge Pack: {manifest.get('businessId', 'unknown')}",
        "",
        f"Generated: {pack['generatedAt']}",
        f"Source root: `{manifest.get('sourceRoot', '')}`",
        "",
        "## Summary",
        "",
        f"- Sources scanned: {summary['sourceCount']}",
        f"- Sources readable: {summary['okCount']}",
        f"- Errors: {summary['errorCount']}",
        f"- Missing: {summary['missingCount']}",
        f"- Document-like sources: {summary['documentCount']}",
        f"- Media sources: {summary['mediaSourceCount']}",
        "",
        "## Extracted Text Sample",
        "",
        summary["combinedTextExcerpt"] or "No text extracted yet.",
        "",
        "## Source Details",
        ""
    ]

    for result in pack["sources"]:
        lines.extend([
            f"### {Path(result['path']).name}",
            "",
            f"- Label: `{result['label']}`",
            f"- Kind: `{result['kind']}`",
            f"- Status: `{result['status']}`",
            f"- Path: `{result['path']}`"
        ])
        if result.get("error"):
            lines.append(f"- Error: {result['error']}")
        if result.get("summary"):
            lines.append("- Summary:")
            for key, value in result["summary"].items():
                if key == "sampleFiles":
                    lines.append(f"  - {key}: {len(value)} sampled")
                elif key == "sheets":
                    lines.append(f"  - {key}: {len(value)} sheet(s)")
                elif key == "tablesSample":
                    lines.append(f"  - {key}: {len(value)} sampled")
                else:
                    lines.append(f"  - {key}: {value}")
        lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Import business source material into a knowledge pack.")
    parser.add_argument("manifest", help="Source manifest YAML path.")
    parser.add_argument("--exclude", action="append", default=[], help="Text or stale reference to redact from extracted excerpts.")
    args = parser.parse_args()

    manifest_path = (PROJECT_ROOT / args.manifest).resolve()
    manifest = parse_manifest(manifest_path)
    source_root = Path(str(manifest["sourceRoot"]))
    business_id = str(manifest.get("businessId", "unknown-business"))

    results = [
        inspect_file(label, resolve_source(source_root, value), args.exclude)
        for label, value in collect_sources(manifest)
    ]

    pack = {
        "businessId": business_id,
        "generatedAt": datetime.now(timezone.utc).isoformat(),
        "manifestPath": str(manifest_path),
        "exclusionsApplied": args.exclude,
        "manifest": manifest,
        "summary": summarize_knowledge_pack(results),
        "sources": [result.__dict__ for result in results]
    }

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    json_path = REPORTS_DIR / f"{business_id}-knowledge-pack.json"
    markdown_path = REPORTS_DIR / f"{business_id}-knowledge-pack.md"
    json_path.write_text(json.dumps(pack, indent=2, ensure_ascii=False), encoding="utf-8")
    write_markdown(pack, markdown_path)

    print(f"Wrote {json_path}")
    print(f"Wrote {markdown_path}")
    print(json.dumps(pack["summary"], indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
